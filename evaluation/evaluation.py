from __future__ import annotations
import csv
import json
import math
import re
import sys
from collections import defaultdict
from dataclasses import dataclass
from pathlib import Path
from statistics import mean
from typing import Any, Iterable

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from retrieve_context import RetrievedChunk, retrieve_context


GROUND_TRUTH_PATH = Path(__file__).with_name("ground_truth.docx")
RESULTS_DIR = Path(__file__).with_name("results")
K_VALUES = (1, 3, 5)
CANDIDATE_TOP_K = 60
REQUIRED_COLUMNS = {
    "case_id",
    "split",
    "question",
    "relevant_targets",
    "validation_status",
}
ACCEPTANCE_TARGETS = {
    "hit_rate_at_1": 0.75,
    "hit_rate_at_3": 0.90,
    "hit_rate_at_5": 0.95,
    "recall_at_5": 0.95,
    "mrr": 0.82,
    "ndcg_at_5": 0.88,
}


@dataclass(frozen=True)
class RelevantTarget:
    book_keyword: str
    hierarchy_keyword: str
    page_number: int | None = None


RelevanceGroup = tuple[RelevantTarget, ...]


def normalize(value: object) -> str:
    text = str(value or "").lower()
    text = (
        text.replace("\u0623", "\u0627")
        .replace("\u0625", "\u0627")
        .replace("\u0622", "\u0627")
        .replace("\u0649", "\u064a")
        .replace("\u0624", "\u0648")
        .replace("\u0626", "\u064a")
    )
    return " ".join(re.findall(r"[a-z0-9\u0600-\u06ff]+", text))


def normalize_header(value: object) -> str:
    return normalize(value).replace(" ", "_")


def parse_target(raw_target: str) -> RelevantTarget:
    parts = [part.strip() for part in str(raw_target or "").split("::")]
    while len(parts) < 3:
        parts.append("")
    page_number: int | None = None
    if parts[2]:
        try:
            page_number = int(parts[2])
        except ValueError:
            page_number = None
    return RelevantTarget(
        book_keyword=normalize(parts[0]),
        hierarchy_keyword=normalize(parts[1]),
        page_number=page_number,
    )


def parse_relevance_groups(value: str) -> list[RelevanceGroup]:
    groups: list[RelevanceGroup] = []
    for raw_group in str(value or "").split("&&"):
        alternatives = tuple(
            parse_target(raw_target)
            for raw_target in raw_group.split("||")
            if raw_target.strip()
        )
        if alternatives:
            groups.append(alternatives)
    return groups


def chunk_matches_target(chunk: RetrievedChunk, target: RelevantTarget) -> bool:
    book = normalize(chunk.book_title)
    hierarchy = normalize(
        " ".join(
            [
                chunk.hierarchy_path,
                chunk.chapter_title,
                chunk.section_title,
                chunk.part_title,
            ]
        )
    )
    if target.book_keyword and target.book_keyword not in book:
        return False
    if target.hierarchy_keyword and target.hierarchy_keyword not in hierarchy:
        return False
    if target.page_number is not None and int(chunk.page_number) != target.page_number:
        return False
    return True


def matching_groups(chunk: RetrievedChunk, groups: list[RelevanceGroup]) -> set[int]:
    return {
        group_index
        for group_index, alternatives in enumerate(groups)
        if any(chunk_matches_target(chunk, target) for target in alternatives)
    }


def relevance_vector(
    chunks: list[RetrievedChunk],
    groups: list[RelevanceGroup],
) -> list[int]:
    matched: set[int] = set()
    vector: list[int] = []
    for chunk in chunks:
        new_groups = matching_groups(chunk, groups) - matched
        vector.append(int(bool(new_groups)))
        matched.update(new_groups)
    return vector


def group_recall_at_k(
    chunks: list[RetrievedChunk],
    groups: list[RelevanceGroup],
    k: int,
) -> float:
    if not groups:
        return 0.0
    matched: set[int] = set()
    for chunk in chunks[:k]:
        matched.update(matching_groups(chunk, groups))
    return len(matched) / len(groups)


def precision_at_k(vector: list[int], k: int) -> float:
    return sum(vector[:k]) / k if k > 0 else 0.0


def hit_rate_at_k(vector: list[int], k: int) -> float:
    return float(any(vector[:k]))


def reciprocal_rank(vector: list[int]) -> float:
    for rank, relevant in enumerate(vector, start=1):
        if relevant:
            return 1.0 / rank
    return 0.0


def ndcg_at_k(vector: list[int], relevant_group_count: int, k: int) -> float:
    dcg = sum(
        relevance / math.log2(rank + 1)
        for rank, relevance in enumerate(vector[:k], start=1)
    )
    ideal_count = min(relevant_group_count, k)
    idcg = sum(1.0 / math.log2(rank + 1) for rank in range(1, ideal_count + 1))
    return dcg / idcg if idcg else 0.0


def load_ground_truth(path: Path = GROUND_TRUTH_PATH) -> list[dict[str, str]]:
    if not path.exists():
        raise FileNotFoundError(f"Ground-truth file was not found: {path}")
    document = Document(str(path))
    if not document.tables:
        raise RuntimeError("ground_truth.docx must contain a table.")
    table = document.tables[0]
    headers = [normalize_header(cell.text) for cell in table.rows[0].cells]
    missing = REQUIRED_COLUMNS - set(headers)
    if missing:
        raise RuntimeError("Missing ground-truth columns: " + ", ".join(sorted(missing)))

    rows: list[dict[str, str]] = []
    for table_row in table.rows[1:]:
        values = [cell.text.strip() for cell in table_row.cells]
        row = {
            header: values[index] if index < len(values) else ""
            for index, header in enumerate(headers)
        }
        if any(row.values()) and normalize(row.get("validation_status")) == "approved":
            rows.append(row)

    if not rows:
        raise RuntimeError("No Approved ground-truth rows were found.")

    case_ids = [row["case_id"] for row in rows]
    if len(case_ids) != len(set(case_ids)):
        raise RuntimeError("Duplicate case_id values were detected.")
    questions = [normalize(row["question"]) for row in rows]
    if len(questions) != len(set(questions)):
        raise RuntimeError("Duplicate ground-truth questions were detected.")
    infant_markers = ("infant", "breastfeeding", "formula feeding", "\u0631\u0636\u064a\u0639", "\u0631\u0636\u0627\u0639\u0647")
    invalid = [
        row["case_id"]
        for row in rows
        if any(marker in normalize(row["question"]) for marker in infant_markers)
    ]
    if invalid:
        raise RuntimeError("Infant-scope questions are not allowed: " + ", ".join(invalid))
    return rows


def evaluate_ranked_list(
    chunks: list[RetrievedChunk],
    groups: list[RelevanceGroup],
    prefix: str,
) -> dict[str, Any]:
    vector = relevance_vector(chunks, groups)
    result: dict[str, Any] = {
        f"{prefix}first_relevant_rank": vector.index(1) + 1 if 1 in vector else "",
        f"{prefix}mrr": round(reciprocal_rank(vector), 6),
    }
    for k in K_VALUES:
        result[f"{prefix}precision_at_{k}"] = round(precision_at_k(vector, k), 6)
        result[f"{prefix}recall_at_{k}"] = round(group_recall_at_k(chunks, groups, k), 6)
        result[f"{prefix}hit_rate_at_{k}"] = round(hit_rate_at_k(vector, k), 6)
        result[f"{prefix}ndcg_at_{k}"] = round(ndcg_at_k(vector, len(groups), k), 6)
    return result


def describe_chunks(chunks: Iterable[RetrievedChunk]) -> str:
    return " || ".join(
        f"{chunk.book_title} > {chunk.hierarchy_path} > page {chunk.page_number}"
        for chunk in chunks
    )


def evaluate_case(row: dict[str, str]) -> dict[str, Any]:
    case_id = row.get("case_id", "").strip()
    question = row.get("question", "").strip()
    groups = parse_relevance_groups(row.get("relevant_targets", ""))
    if not case_id or not question or not groups:
        raise ValueError(f"Invalid ground-truth row: {case_id or '<missing id>'}")

    retrieval = retrieve_context(
        question=question,
        top_k=CANDIDATE_TOP_K,
        final_k=max(K_VALUES),
    )
    final_chunks = list(retrieval.get("chunks", []))
    dense_chunks = list(retrieval.get("dense_chunks", final_chunks))
    hybrid_chunks = list(retrieval.get("hybrid_chunks", final_chunks))
    reranked_chunks = list(retrieval.get("reranked_chunks", final_chunks))

    result: dict[str, Any] = {
        "case_id": case_id,
        "split": row.get("split", "test").strip().lower(),
        "question": question,
        "language": row.get("language", ""),
        "category": row.get("category", ""),
        "difficulty": row.get("difficulty", ""),
        "age_scope": row.get("age_scope", ""),
        "relevance_group_count": len(groups),
        "accepted_alternative_count": sum(len(group) for group in groups),
        "rewritten_query": retrieval.get("rewritten_query", ""),
        "expanded_query": retrieval.get("expanded_query", ""),
        "retrieved_chapters": describe_chunks(final_chunks),
    }
    result.update(evaluate_ranked_list(dense_chunks, groups, "dense_"))
    result.update(evaluate_ranked_list(hybrid_chunks, groups, "hybrid_"))
    result.update(evaluate_ranked_list(reranked_chunks, groups, "reranked_"))
    result.update(evaluate_ranked_list(final_chunks, groups, ""))
    return result


def mean_metric(rows: list[dict[str, Any]], key: str) -> float:
    return round(mean(float(row[key]) for row in rows), 6) if rows else 0.0


def aggregate_rows(rows: list[dict[str, Any]]) -> dict[str, Any]:
    summary: dict[str, Any] = {"cases": len(rows), "mrr": mean_metric(rows, "mrr")}
    for k in K_VALUES:
        summary[f"precision_at_{k}"] = mean_metric(rows, f"precision_at_{k}")
        summary[f"recall_at_{k}"] = mean_metric(rows, f"recall_at_{k}")
        summary[f"hit_rate_at_{k}"] = mean_metric(rows, f"hit_rate_at_{k}")
        summary[f"ndcg_at_{k}"] = mean_metric(rows, f"ndcg_at_{k}")
    summary["stage_hit_at_5"] = {
        "dense": mean_metric(rows, "dense_hit_rate_at_5"),
        "hybrid": mean_metric(rows, "hybrid_hit_rate_at_5"),
        "reranked": mean_metric(rows, "reranked_hit_rate_at_5"),
        "final": mean_metric(rows, "hit_rate_at_5"),
    }
    return summary


def grouped_breakdown(rows: list[dict[str, Any]], field: str) -> dict[str, Any]:
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for row in rows:
        grouped[str(row.get(field, "Unknown") or "Unknown")].append(row)
    return {key: aggregate_rows(value) for key, value in sorted(grouped.items())}


def build_summary(
    successful: list[dict[str, Any]],
    failed: list[dict[str, Any]],
) -> dict[str, Any]:
    dev_rows = [row for row in successful if row.get("split") == "dev"]
    test_rows = [row for row in successful if row.get("split") == "test"]
    primary = aggregate_rows(test_rows)
    acceptance = {
        metric: {
            "actual": float(primary.get(metric, 0.0)),
            "target": target,
            "passed": float(primary.get(metric, 0.0)) >= target,
        }
        for metric, target in ACCEPTANCE_TARGETS.items()
    }
    return {
        "total_cases": len(successful) + len(failed),
        "successful_cases": len(successful),
        "failed_cases": len(failed),
        "dev": aggregate_rows(dev_rows),
        "test": primary,
        "all": aggregate_rows(successful),
        "test_by_language": grouped_breakdown(test_rows, "language"),
        "test_by_category": grouped_breakdown(test_rows, "category"),
        "test_by_difficulty": grouped_breakdown(test_rows, "difficulty"),
        "acceptance": acceptance,
    }


def save_csv(path: Path, rows: list[dict[str, Any]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if not rows:
        path.write_text("", encoding="utf-8")
        return
    columns = sorted({key for row in rows for key in row})
    with path.open("w", newline="", encoding="utf-8-sig") as file:
        writer = csv.DictWriter(file, fieldnames=columns)
        writer.writeheader()
        writer.writerows(rows)


def save_results(
    successful: list[dict[str, Any]],
    failed: list[dict[str, Any]],
    summary: dict[str, Any],
) -> None:
    RESULTS_DIR.mkdir(parents=True, exist_ok=True)
    save_csv(RESULTS_DIR / "retrieval_results.csv", successful)
    save_csv(RESULTS_DIR / "evaluation_errors.csv", failed)
    (RESULTS_DIR / "evaluation_summary.json").write_text(
        json.dumps(summary, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )


def print_split(label: str, summary: dict[str, Any]) -> None:
    print(f"\n{label}")
    print("-" * 72)
    print(f"Cases: {summary['cases']}")
    print(f"MRR: {summary['mrr']:.4f}")
    for k in K_VALUES:
        print(
            f"K={k}: Hit={summary[f'hit_rate_at_{k}']:.4f} | "
            f"Recall={summary[f'recall_at_{k}']:.4f} | "
            f"nDCG={summary[f'ndcg_at_{k}']:.4f} | "
            f"Chapter Precision={summary[f'precision_at_{k}']:.4f}"
        )
    stages = summary["stage_hit_at_5"]
    print(
        "Stage Hit@5: "
        f"Dense={stages['dense']:.4f}, Hybrid={stages['hybrid']:.4f}, "
        f"Reranked={stages['reranked']:.4f}, Final={stages['final']:.4f}"
    )


def print_summary(summary: dict[str, Any]) -> None:
    print("\nRETRIEVAL EVALUATION SUMMARY")
    print("=" * 72)
    print(f"Total cases: {summary['total_cases']}")
    print(f"Successful cases: {summary['successful_cases']}")
    print(f"Failed cases: {summary['failed_cases']}")
    print_split("DEV SET - tuning only", summary["dev"])
    print_split("TEST SET - presentation metrics", summary["test"])
    print("\nTEST ACCEPTANCE TARGETS")
    print("-" * 72)
    for metric, result in summary["acceptance"].items():
        status = "PASS" if result["passed"] else "NEEDS TUNING"
        print(
            f"{metric}: {result['actual']:.4f} / {result['target']:.4f} -> {status}"
        )
    print(
        "\nChapter Precision@K is descriptive only when a question has one "
        "relevance group; Hit Rate, group Recall, MRR, and nDCG are the primary metrics."
    )


def main() -> None:
    rows = load_ground_truth()
    successful: list[dict[str, Any]] = []
    failed: list[dict[str, Any]] = []

    for index, row in enumerate(rows, start=1):
        case_id = row.get("case_id", "")
        print(f"[{index}/{len(rows)}] Evaluating {case_id}...")
        try:
            successful.append(evaluate_case(row))
        except Exception as error:
            failed.append(
                {
                    "case_id": case_id,
                    "question": row.get("question", ""),
                    "error": f"{type(error).__name__}: {error}",
                }
            )

    if not successful:
        raise RuntimeError("All evaluation cases failed.")

    summary = build_summary(successful, failed)
    print_summary(summary)
    save_results(successful, failed, summary)
    print("\nResults saved inside evaluation/results.")


if __name__ == "__main__":
    main()