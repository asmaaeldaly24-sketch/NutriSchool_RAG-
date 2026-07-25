from __future__ import annotations

import json
import py_compile
import re
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parent
PYTHON_FILES = (
    "book_catalog.py",
    "documents.py",
    "preprocessing.py",
    "chunking.py",
    "embeddings.py",
    "vector_store.py",
    "build_index.py",
    "retrieve_context.py",
    "prompting.py",
    "rag_pipeline.py",
    "evaluation/evaluation.py",
    "evaluation/tune_retrieval.py",
)


def normalize(value: object) -> str:
    return " ".join(re.findall(r"[a-z0-9]+", str(value or "").lower()))


def read_ground_truth() -> list[dict[str, str]]:
    path = PROJECT_ROOT / "evaluation" / "ground_truth.docx"
    document = Document(str(path))
    if not document.tables:
        raise RuntimeError("ground_truth.docx does not contain a table.")
    table = document.tables[0]
    headers = [normalize(cell.text).replace(" ", "_") for cell in table.rows[0].cells]
    rows: list[dict[str, str]] = []
    for table_row in table.rows[1:]:
        values = [cell.text.strip() for cell in table_row.cells]
        row = {
            header: values[index] if index < len(values) else ""
            for index, header in enumerate(headers)
        }
        if any(row.values()):
            rows.append(row)
    return rows


def main() -> None:
    for relative in PYTHON_FILES:
        path = PROJECT_ROOT / relative
        if not path.is_file():
            raise FileNotFoundError(f"Missing required file: {relative}")
        py_compile.compile(str(path), doraise=True)

    rows = read_ground_truth()
    dev_count = sum(row.get("split", "").lower() == "dev" for row in rows)
    test_count = sum(row.get("split", "").lower() == "test" for row in rows)
    if dev_count != 6 or test_count != 30:
        raise RuntimeError(
            f"Expected 6 dev and 30 test cases; found {dev_count} dev and {test_count} test."
        )
    case_ids = [row.get("case_id", "") for row in rows]
    questions = [row.get("question", "").strip().lower() for row in rows]
    if len(case_ids) != len(set(case_ids)):
        raise RuntimeError("Duplicate case IDs were detected.")
    if len(questions) != len(set(questions)):
        raise RuntimeError("Duplicate questions were detected.")

    manifest_path = PROJECT_ROOT / "index_manifest.json"
    manifest = None
    if manifest_path.is_file():
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        if manifest.get("audit", {}).get("suspicious_chunk_count") != 0:
            raise RuntimeError("The index manifest reports suspicious chunks.")

    print("PROJECT VALIDATION PASSED")
    print(f"Python files: {len(PYTHON_FILES)}")
    print(f"Ground truth: {dev_count} dev + {test_count} test")
    print(f"Index manifest present: {manifest is not None}")


if __name__ == "__main__":
    main()